# --- Import Libraries ---
# We need 'Document' from the 'docx' library to read the file.
from docx import Document
# We import 'os' to help find the file.
import os

# --- Define the Function ---
# A new function to load our Word document.
def load_docx_data(file_path):
    """
    This function loads text from a .docx file (the EU AI Act).
    It reads paragraph by paragraph and joins them.
    """
    
    print(f"--- 2. Loading DOCX file: {file_path} ---")

    # --- Pre-processing Explanation ---
    # DECISION: A Word document is structured as a list of
    # paragraphs. We will read each paragraph one by one and
    # join them together. We will put two "newline" characters ('\n\n')
    # between them. This is like pressing 'Enter' twice - it leaves
    # a blank line, which keeps the document's structure clear for the AI.

    try:
        # This opens the Word document file.
        document = Document(file_path)
        
        # Create an empty list to store all the text.
        text_paragraphs = []
        
        # This is a 'for' loop. It repeats for every paragraph
        # in the 'document.paragraphs' list.
        for para in document.paragraphs:
            # We add the text from the paragraph (para.text)
            # to our 'text_paragraphs' list.
            text_paragraphs.append(para.text)
            
        # This joins everything in the list into one big string.
        # The '\n\n' separates each paragraph with a blank line.
        full_text = "\n\n".join(text_paragraphs)
        
        print("--- ...DOCX loaded and joined successfully. ---")
        
        # Send the final text back.
        return full_text

    # Catch errors if the file is missing or corrupt.
    except FileNotFoundError:
        print(f"!!! ERROR: The file '{file_path}' was not found. !!!")
        return None
    except Exception as e:
        print(f"!!! ERROR: An error occurred: {e} !!!")
        return None

# --- Self-Test Block ---
# This block is for testing this file by itself.
if __name__ == "__main__":
    print("="*70)
    print("TESTING load_doc.py")
    print("="*70)
    
    # DECISION: Try multiple possible paths for the test file
    # because it might be in the root folder OR in the data/documents/ folder
    # This makes the test more robust and flexible
    
    test_file = 'EU AI Act Doc (1) (3).docx'
    
    # Check different possible locations
    if os.path.exists(test_file):
        test_path = test_file
    elif os.path.exists(os.path.join('data/documents', test_file)):
        test_path = os.path.join('data/documents', test_file)
    elif os.path.exists(os.path.join('../../data/documents', test_file)):
        test_path = os.path.join('../../data/documents', test_file)
    else:
        print(f"!!! Could not find {test_file} in data/documents/ folder !!!")
        test_path = None
    
    # Run the function if we found the file
    if test_path:
        content = load_docx_data(test_path)
        
        # If it worked, print a sample.
        if content:
            print("\n" + "="*70)
            print("PREVIEW: First 500 characters")
            print("="*70)
            print(content[:500])
            print("\n...")
            print(f"\nTotal length: {len(content)} characters")